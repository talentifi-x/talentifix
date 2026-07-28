# -*- coding: utf-8 -*-
"""Generate the TalentiFi-X end-to-end website audit report (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette ----------
PRIMARY   = "0000FF"
SECONDARY = "00B3B8"   # darkened cyan for contrast on white
DARK      = "1E1E24"
GREY      = "5D6D7E"
SEV = {
    "CRITICAL": "C00000",
    "HIGH":     "D35400",
    "MEDIUM":   "B7950B",
    "LOW":      "2E86C1",
    "POSITIVE": "1E8449",
}
SEV_BG = {
    "CRITICAL": "F8D7DA",
    "HIGH":     "FCE4D6",
    "MEDIUM":   "FFF2CC",
    "LOW":      "DDEBF7",
    "POSITIVE": "D5F5E3",
}

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)

for h, sz, col in (("Heading 1", 17, PRIMARY), ("Heading 2", 13.5, DARK), ("Heading 3", 11.5, PRIMARY)):
    st = doc.styles[h]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor.from_string(col)
    st.font.bold = True

# ---------- helpers ----------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
    elif color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p

def para(text="", size=10.5, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
    return p

def bullet(text, lead=None, lead_color=None, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead + " ")
        r.font.bold = True
        r.font.size = Pt(size)
        if lead_color:
            r.font.color.rgb = RGBColor.from_string(lead_color)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    return p

def sev_run(p, sev):
    r = p.add_run(f" {sev} ")
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
    # shading on run
    rpr = r._element.get_or_add_rPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), SEV[sev])
    rpr.append(shd)
    return r

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "BFC9CA")
    pbdr.append(bottom); pPr.append(pbdr)

def page_break():
    doc.add_page_break()

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(2):
    doc.add_paragraph()
p = para("WEBSITE AUDIT REPORT", size=30, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("End-to-End Technical, SEO, Performance, Accessibility & Security Review",
     size=13, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
para("TalentiFi-X", size=22, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("https://www.talentifix.com", size=12, color=SECONDARY, align=WD_ALIGN_PARAGRAPH.CENTER, after=40)

# cover meta table
ct = doc.add_table(rows=5, cols=2)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
ct.style = "Light List Accent 1"
cover_meta = [
    ("Report Date", "18 June 2026"),
    ("Audited Environment", "Production (Vercel) + current source tree (branch: karan)"),
    ("Scope", "Full site — SEO, Security, Performance, Accessibility, Code Quality"),
    ("Methodology", "Static source review, live HTTP/SSR inspection, Lighthouse (mobile + desktop)"),
    ("Confidentiality", "Internal — contains security-sensitive findings"),
]
for i, (k, v) in enumerate(cover_meta):
    set_cell_text(ct.rows[i].cells[0], k, bold=True, size=10)
    set_cell_text(ct.rows[i].cells[1], v, size=10)
page_break()

# ============================================================
# TABLE OF CONTENTS (field — update with F9 in Word)
# ============================================================
doc.add_heading("Contents", level=1)
para("This document contains an automatic table of contents. To populate it in Microsoft Word: "
     "click anywhere in the area below and press F9 (or right-click → Update Field).",
     size=9, italic=True, color=GREY, after=10)
p = doc.add_paragraph()
run = p.add_run()
fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
t = OxmlElement("w:t"); t.text = "Right-click and select “Update Field” to build the table of contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
for el in (fldChar, instrText, fldChar2, t, fldChar3):
    run._r.append(el)
page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading("1. Executive Summary", level=1)
para("TalentiFi-X is a well-designed, modern marketing site (Next.js 16 App Router, React 19, "
     "Sanity CMS, Tailwind v4) hosted on Vercel. The codebase is clean and the desktop experience is "
     "fast. However, the audit uncovered one architecture-level defect that severely undermines the "
     "site’s search visibility, a confirmed credential exposure in version-control history, and a "
     "fragmented domain/canonical setup. Together these mean the site is currently underperforming its "
     "potential for organic discovery despite strong underlying content and design.",
     after=8)

para("Headline verdict", bold=True, size=11, color=PRIMARY, after=2)
bullet("The entire production site renders client-side with an EMPTY server-side HTML shell "
       "(every page returns the Next.js “BAILOUT_TO_CLIENT_SIDE_RENDERING” marker, with zero headings, "
       "links, images, or body text in the server response). Search engines and social/AI crawlers that do "
       "not fully execute JavaScript effectively see blank pages. This is the most important issue in the report.",
       lead="CRITICAL —", lead_color=SEV["CRITICAL"])
bullet("Live SMTP and Sanity write credentials were committed to git history and pushed to GitHub. "
       "Keys have been rotated (confirmed by owner), but the history still needs to be purged.",
       lead="CRITICAL —", lead_color=SEV["CRITICAL"])
bullet("Domain/canonical signals are fragmented across three different host spellings; the XML sitemap "
       "lists URLs that all 301-redirect, and blog canonicals point to a host that redirects elsewhere.",
       lead="CRITICAL —", lead_color=SEV["CRITICAL"])
bullet("No structured data, no social-share (OpenGraph) image, sitemap omits all job pages, and consent "
       "is not enforced before analytics cookies load (GDPR risk).", lead="HIGH —", lead_color=SEV["HIGH"])
bullet("112 MB of unoptimised source images (an 8.5 MB hero) drive a poor mobile experience "
       "(Lighthouse mobile Performance 47, LCP 6.4 s) — although desktop scores 97.",
       lead="HIGH —", lead_color=SEV["HIGH"])

# severity count table
para("Findings by severity", bold=True, size=11, color=PRIMARY, before=8, after=4)
counts = [("Critical", 3, SEV["CRITICAL"]),
          ("High", 12, SEV["HIGH"]),
          ("Medium", 19, SEV["MEDIUM"]),
          ("Low", 12, SEV["LOW"])]
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for c, lbl in zip(hdr, ["Critical", "High", "Medium", "Low"]):
    pass
for i, (lbl, n, col) in enumerate(counts):
    set_cell_text(hdr[i], lbl, bold=True, white=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    shade(hdr[i], col)
row = tbl.add_row().cells
for i, (lbl, n, col) in enumerate(counts):
    set_cell_text(row[i], str(n), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=col)
para("46 distinct findings in total, plus a list of strengths the site already does well (Section 6).",
     size=9, italic=True, color=GREY, before=4)

# scorecard
para("Live Lighthouse scorecard (production, 18 Jun 2026)", bold=True, size=11, color=PRIMARY, before=10, after=4)
sc = doc.add_table(rows=1, cols=6)
sc.style = "Table Grid"
heads = ["Profile", "Performance", "Accessibility", "Best Practices", "SEO*", "LCP"]
for c, h in zip(sc.rows[0].cells, heads):
    set_cell_text(c, h, bold=True, white=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
    shade(c, PRIMARY)
def score_cell(cells, idx, val, good_thresh=90):
    col = SEV["POSITIVE"] if isinstance(val, int) and val >= good_thresh else (SEV["MEDIUM"] if isinstance(val, int) and val >= 50 else SEV["CRITICAL"])
    set_cell_text(cells[idx], str(val), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=col if isinstance(val, int) else DARK)
for prof, perf, a11y, bp, seo, lcp in [("Mobile", 47, 96, 100, 100, "6.4 s"), ("Desktop", 97, 96, 100, 100, "1.0 s")]:
    cells = sc.add_row().cells
    set_cell_text(cells[0], prof, bold=True, size=10)
    score_cell(cells, 1, perf); score_cell(cells, 2, a11y); score_cell(cells, 3, bp); score_cell(cells, 4, seo)
    set_cell_text(cells[5], lcp, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
                  color=SEV["CRITICAL"] if "6" in lcp else SEV["POSITIVE"])
para("*Lighthouse’s automated SEO score (100) only checks basic on-page tags. It does NOT detect the "
     "client-side-rendering bailout, the canonical/domain conflicts, the missing structured data, or the "
     "absent social image — all of which are covered in this report. Treat the 100 as “no trivial errors,” "
     "not “SEO is healthy.”", size=9, italic=True, color=GREY, before=4)
page_break()

# ============================================================
# METHODOLOGY
# ============================================================
doc.add_heading("2. Methodology & Scope", level=1)
para("The audit combined three complementary techniques:", after=4)
bullet("Static source review of the full codebase (Next.js App Router pages, React components, API routes, "
       "Sanity integration, configuration) across five dimensions: SEO, Security, Performance, Accessibility, "
       "and Code Quality.")
bullet("Live production inspection of https://www.talentifix.com — HTTP headers, redirect topology, "
       "server-rendered HTML (SSR) for every primary route, robots.txt, sitemap.xml, asset transfer sizes, "
       "compression, and security headers.")
bullet("Lighthouse audits (v12, headless Chrome) for mobile and desktop, capturing Core Web Vitals and "
       "diagnostics under lab conditions.")
para("What was NOT in scope / not available:", bold=True, before=8, after=2)
bullet("Analytics and Search Console data (Google Analytics, GSC, Clarity) — excluded at owner’s request; "
       "traffic, indexation-coverage and conversion analysis are therefore not included.")
bullet("Field (real-user) Core Web Vitals via CrUX/PageSpeed API — the keyless quota was exhausted; figures "
       "here are Lighthouse lab data.")
bullet("Penetration testing / live exploitation — security findings are based on code and response inspection.")

doc.add_heading("Severity model", level=2)
for s, desc in [
    ("CRITICAL", "Breaks a core capability (search visibility, security, or data integrity). Fix immediately."),
    ("HIGH", "Significant business/SEO/UX impact or a real risk; schedule in the current sprint."),
    ("MEDIUM", "Meaningful quality, compliance, or maintainability issue; plan into upcoming work."),
    ("LOW", "Minor polish, hygiene, or best-practice gap; address opportunistically."),
]:
    p = para("", after=3)
    sev_run(p, s)
    r = p.add_run("  " + desc); r.font.size = Pt(10)
page_break()

# ============================================================
# DETAILED FINDINGS — data model
# ============================================================
# Each finding: (id, severity, title, area, evidence, impact, fix)
FINDINGS = [
 # ---------------- CRITICAL ----------------
 ("C-01", "CRITICAL", "Entire site renders client-side with an empty server HTML shell (no SSR content)",
  "SEO / Architecture / Performance",
  "Live: every audited route (/, /about, /solutions, /blog, /contact, /jobs, blog post) returns "
  "`<template data-dgst=\"BAILOUT_TO_CLIENT_SIDE_RENDERING\">` and ZERO rendered <h1–h6>, <a>, <img>, or <p> "
  "tags in the server response (only the <head> metadata + a ~83 KB React payload). Root cause in source: "
  "src/app/layout.tsx wraps the whole app in a single <Suspense> around <Refine options={{syncWithLocation:true}}>; "
  "Refine’s router uses useSearchParams(), which forces Next.js to de-opt the entire tree to client-side rendering.",
  "Crawlers, social scrapers and AI bots that don’t fully execute JavaScript see blank pages — all headings, "
  "copy, internal links and structured data are invisible to them. Google can render JS but does so on a slower, "
  "less reliable second pass. This also delays first paint (contributes to the 6.4 s mobile LCP). It silently "
  "negates almost every other SEO effort on the site.",
  "Remove Refine entirely (it is unused — see C-Q findings; its data provider points at a demo API). If Refine "
  "must stay, do not wrap page content in it and isolate useSearchParams() behind a narrow Suspense boundary so "
  "pages server-render. After the fix, re-verify that `curl https://www.talentifix.com/` returns real <h1>/<p>/<a> "
  "markup and that the BAILOUT marker is gone."),

 ("C-02", "CRITICAL", "Live SMTP password and Sanity write token committed to git history (pushed to GitHub)",
  "Security",
  "A .env file containing BREVO_SMTP_PASS and a Sanity `sk` write token was committed across ~7 historical "
  "commits (earliest 1a4ac80) and pushed to the remote on branches master and karan. The file was later removed "
  "from the working tree but remains fully recoverable from history (e.g. `git show <commit>:.env`).",
  "Anyone with repository access can recover the credentials and (a) send mail through the company Brevo account "
  "(spam/phishing, sender-reputation damage) and (b) read, modify or delete all CMS content via the Sanity write token.",
  "Keys have been ROTATED (confirmed by owner) — good. Remaining action: purge the file from history with "
  "git-filter-repo or BFG, force-push all branches, and have collaborators re-clone. Keep secrets only in "
  ".env.local (already git-ignored) and Vercel’s encrypted env store. Treat the old keys as permanently burned."),

 ("C-03", "CRITICAL", "Fragmented canonical domain — sitemap and blog canonicals point at redirecting hosts",
  "SEO",
  "Live redirect map: content serves ONLY at https://www.talentifix.com (200). All of www.talentifi-x.com "
  "(the host the owner believed canonical), talentifi-x.com, and talentifix.com issue 301s to it. BUT: the XML "
  "sitemap lists https://talentifix.com/… URLs (no www) — every entry 301-redirects; robots.txt references the "
  "same redirecting sitemap host; and blog post canonicals are hard-coded to https://www.talentifi-x.com/blog/… "
  "(src/app/blog/[slug]/page.tsx) — a third spelling that redirects to a different host. Homepage and main pages "
  "emit no canonical tag at all.",
  "Conflicting/again redirecting canonical signals dilute ranking authority, waste crawl budget on redirect hops, "
  "and can cause Google to ignore the declared canonicals. A sitemap full of non-200 URLs is a recognised quality issue.",
  "Pick ONE canonical origin: https://www.talentifix.com. Set NEXT_PUBLIC_SITE_URL to it in Vercel and derive every "
  "URL from that single value (sitemap, robots, all canonicals, internal absolute links). Remove the hard-coded "
  "talentifi-x.com in blog/[slug]. Add a self-referencing canonical to every page via metadataBase + alternates."),

 # ---------------- HIGH ----------------
 ("H-01", "HIGH", "No structured data (JSON-LD) anywhere on the site",
  "SEO",
  "Live: 0 `application/ld+json` blocks on every page. Missing Organization, JobPosting (jobs have all required "
  "fields), Article/BlogPosting, FAQPage (blog posts render FAQs), and BreadcrumbList (jobs render a visual breadcrumb).",
  "Forfeits rich results — most notably Google Jobs eligibility for /jobs/* and FAQ/Article enhancements — and "
  "weakens entity understanding for AI search.",
  "Add JSON-LD: Organization (logo, name, sameAs, both addresses) in the root layout; JobPosting per job page; "
  "Article + FAQPage on blog posts; BreadcrumbList where breadcrumbs appear. (Must be server-rendered — depends on C-01.)"),

 ("H-02", "HIGH", "No OpenGraph / social-share image",
  "SEO / Brand",
  "Live: /opengraph-image, /opengraph-image.png and /og.png all return 404; og:image is absent from page HTML. "
  "No default image is set in root metadata and no per-page images exist except when a Sanity post supplies one.",
  "Links shared on LinkedIn, X, WhatsApp, Slack, etc. render as bare text with no image — poor CTR and brand presentation.",
  "Add a branded 1200×630 default (src/app/opengraph-image.png or openGraph.images in root metadata) plus a Twitter "
  "image; add per-page overrides for key landing pages."),

 ("H-03", "HIGH", "metadataBase is not set — OG/relative URLs cannot resolve",
  "SEO",
  "No `metadataBase` in src/app/layout.tsx. Next.js cannot resolve relative OpenGraph/Twitter image URLs or "
  "openGraph.url to absolute URLs and emits a build warning.",
  "Even once an OG image exists, social crawlers may reject relative URLs; canonical/OG URL generation is unreliable.",
  "Set `metadataBase: new URL(process.env.NEXT_PUBLIC_SITE_URL ?? 'https://www.talentifix.com')` in root metadata."),

 ("H-04", "HIGH", "XML sitemap omits all job pages (and lists redirecting URLs)",
  "SEO",
  "src/app/sitemap.ts enumerates static routes + Sanity blog slugs only — no /jobs and no /jobs/[slug], although "
  "getAllSanityJobSlugs already exists. Combined with C-03, every listed URL also 301-redirects.",
  "Job detail pages (high commercial intent, Jobs-rich-result candidates) are harder to discover and index.",
  "Add /jobs and map getAllSanityJobSlugs() into /jobs/[slug] entries (try/catch like blog). Emit final-canonical "
  "(www) URLs. Mirror in the HTML sitemap."),

 ("H-05", "HIGH", "No rate limiting, CAPTCHA, or bot protection on any form API",
  "Security",
  "All four POST routes (contact, candidate-registration, job-application, primary-client-contact) have no "
  "throttling, CAPTCHA, or honeypot. Each submit sends TWO emails (internal + auto-reply to the attacker-supplied "
  "address) and runs transporter.verify(). The site is on Vercel’s free plan (no WAF).",
  "Abuse can exhaust the Brevo sending quota / incur cost, use the auto-reply as a spam reflector to arbitrary "
  "victims, and hammer the mail server. Confirmed: Sanity CORS is also currently unrestricted.",
  "Add per-IP rate limiting (e.g. @upstash/ratelimit) and a honeypot or Cloudflare/Turnstile CAPTCHA on all four "
  "routes; remove per-request transporter.verify(); restrict Sanity CORS origins to known hosts."),

 ("H-06", "HIGH", "Cookie consent is cosmetic — analytics load before/without consent (GDPR/ePrivacy)",
  "Security / Compliance",
  "src/app/layout.tsx loads Google Analytics (G-VDENLSNNWP) and Microsoft Clarity unconditionally with "
  "strategy=afterInteractive. src/components/layout/CookieConsent.tsx only writes a localStorage flag and never "
  "gates the scripts; clicking “Reject” has no effect.",
  "Non-essential tracking cookies are set before consent for EU/UK visitors — a GDPR/ePrivacy violation; the banner "
  "currently provides no real choice.",
  "Gate the GA/Clarity <Script> tags on stored consent (mount only when accepted), or implement GA Consent Mode v2 "
  "defaulting to denied and updating on accept. Wire the Reject button to actually suppress loading."),

 ("H-07", "HIGH", "112 MB of unoptimised source images (8.5 MB LCP hero) — poor mobile performance",
  "Performance",
  "public/ holds 112.1 MB of images across 103 files: banner.webp 8.48 MB (the priority-loaded LCP hero), "
  "Next.webp 7.88 MB, contact-form-bg.png 7.88 MB, multiple 6–7 MB blog/solution images, who-polygon.png 4.78 MB, "
  "a 2.21 MB “SVG”. Lighthouse mobile: Performance 47, FCP 4.0 s, LCP 6.4 s, TBT 720 ms (CLS 0).",
  "Although next/image re-encodes on the fly, the giant sources strain the optimiser (cold-start cost) and the hero "
  "dominates mobile LCP. Desktop is fine (LCP 1.0 s) but most recruiting/candidate traffic is mobile.",
  "Pre-compress all sources to web scale before they hit the optimiser (hero ≈150–400 KB at display size). Re-export "
  "blog PNGs to WebP/AVIF; make the 2.2 MB SVG a true vector or a sized raster. Add preconnect for cdn.sanity.io."),

 ("H-08", "HIGH", "Application modal has no focus trap, initial focus, or focus return",
  "Accessibility",
  "src/components/jobs/ApplyModal.tsx sets role=dialog, aria-modal and Escape-to-close (good) but never moves focus "
  "into the dialog on open, does not trap Tab within it (focusable page elements remain reachable behind the backdrop), "
  "and does not return focus to the trigger on close.",
  "Keyboard and screen-reader users can be left navigating the page behind the modal — fails WCAG 2.4.3 / 4.1.2.",
  "On open, move focus to the dialog/first field; trap Tab/Shift+Tab; store and restore document.activeElement on close. "
  "Consider `inert` on the background or a headless dialog primitive."),

 ("H-09", "HIGH", "Form errors and success messages are not announced to assistive tech",
  "Accessibility",
  "Status/error banners in ContactForm, PrimaryClientContactForm, CandidateRegistrationForm and ApplyModal are plain "
  "<div>s with no role=alert / aria-live; inline field errors are not tied to inputs via aria-invalid / aria-describedby.",
  "Screen-reader users who submit an invalid form receive no spoken feedback — fails WCAG 3.3.1 / 4.1.3.",
  "Add role=alert + aria-live=assertive to error containers (role=status for success) and wire field errors with "
  "aria-invalid + aria-describedby."),

 ("H-10", "HIGH", "Missing / mis-ordered H1s; decorative footer <h1> on every page",
  "Accessibility / SEO",
  "In source, Solutions, Contact, Start-Hiring and Join-Our-Network pages have no <h1> (start at <h2>/<h3>); the "
  "footer renders a decorative <h1>“Talentifi-X” on every page; src/app/page.tsx nests a second <main> inside the "
  "layout’s <main>. (Note: live, all pages currently show 0 headings due to C-01 — fix C-01 first, then these.)",
  "Confused heading outline and duplicate/empty H1s weaken both accessibility landmarks and on-page SEO once SSR is restored.",
  "Give each page exactly one meaningful <h1>; demote the footer brand text to a non-heading aria-hidden element; "
  "change the home page wrapper from <main> to <div>."),

 ("H-11", "HIGH", "Sanity client uses useCdn:false — every read bypasses the CDN",
  "Performance",
  "src/sanity/lib/client.ts sets useCdn:false, routing all published-content reads to the live API instead of the "
  "cached CDN. Pages use revalidate=60 ISR, so CDN caching would compose well.",
  "Adds latency and API load to every blog/job render; unnecessary for published reads.",
  "Set useCdn:true for the public read client; use a separate authenticated (useCdn:false) client only for any "
  "preview/draft route."),

 ("H-12", "HIGH", "Sanity Studio and the Refine framework are bundled into the marketing site",
  "Performance / Architecture",
  "The full Sanity Studio (sanity, @sanity/vision) is mounted at /studio and the entire Refine stack "
  "(@refinedev/core, kbar, devtools, nextjs-router, simple-rest) wraps every page in layout.tsx — yet no admin/CRUD "
  "feature uses Refine and its data provider points at api.fake-rest.refine.dev. (Refine is also the root cause of C-01.)",
  "Large, unnecessary dependencies inflate install/build and client JS; Refine actively breaks SSR (C-01).",
  "Remove Refine completely (provider, layout wrapper, 5 packages, the refine dev/build/start scripts). If the "
  "embedded Studio isn’t used by editors, remove it too or host it as a separate deployment."),

 # ---------------- MEDIUM ----------------
 ("M-01", "MEDIUM", "No Content-Security-Policy; HSTS lacks includeSubDomains/preload",
  "Security",
  "Live headers include X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy and HSTS "
  "(max-age=63072000) — good — but there is no Content-Security-Policy, and HSTS omits includeSubDomains; preload. "
  "No CSP is defined in next.config.mjs.",
  "Without CSP there is no defence-in-depth against XSS (compounds M-02); HSTS is not preload-eligible.",
  "Add a CSP allowing self + known third parties (googletagmanager, clarity.ms, cdn.sanity.io, GA) using a nonce/hash "
  "for the inline GA/Clarity scripts; extend HSTS to `includeSubDomains; preload`."),

 ("M-02", "MEDIUM", "PortableText link mark allows javascript: (and other) URI schemes — stored-XSS vector",
  "Security",
  "src/app/blog/[slug]/page.tsx renders link marks as href = value.href ?? '#' with no scheme validation (only a "
  "leading-slash check). The image asset.url is likewise unvalidated.",
  "A CMS editor (or anyone who had obtained the leaked write token) could author href=\"javascript:…\" producing "
  "stored XSS for every blog reader.",
  "Allow only http/https/mailto and leading-slash hrefs; otherwise drop to '#'. Apply the same allow-list to image URLs."),

 ("M-03", "MEDIUM", "Sanity Studio publicly mounted at /studio; CORS unrestricted",
  "Security",
  "The Studio route is reachable by anyone (robots disallow is a crawl hint, not access control). Owner confirms it "
  "is now restricted at the hosting level (good) but Sanity CORS origins are NOT limited, and projectId/apiVersion are public.",
  "Combined with public CORS, the Sanity API is reachable directly from any origin; the admin UI is discoverable.",
  "Confirm the /studio edge restriction; restrict Sanity CORS to known hosts; ensure no anonymous dataset write access."),

 ("M-04", "MEDIUM", "Forced quality={100} on multiple images defeats the optimiser",
  "Performance",
  "HumanLeadSection, SolutionSection, RebuildSection (x2), WhereWeSpecializeSection and TheNextStepSection pass "
  "quality={100}; next.config.mjs even allow-lists qualities [100,75].",
  "Produces the largest possible re-encodes (often 2–4× the bytes of q75) for negligible visual gain.",
  "Use quality={75} (default) or drop the prop; remove 100 from the qualities array to prevent regressions."),

 ("M-05", "MEDIUM", "Self-hosted fonts use TTF, bypass next/font, and reference a non-existent face",
  "Performance",
  "src/styles/global.css declares six @font-face rules loading ~76 KB .ttf files (uncompressed vs WOFF2) with no "
  "font-display; --font-stack-text points at “Stack Sans Text” which has no @font-face (silently falls back to Inter).",
  "Larger font transfers, potential invisible-text flash (FOIT), and a dead font variable.",
  "Convert to WOFF2 and load via next/font/local (auto preload + swap), or at minimum add font-display:swap; define "
  "or remove the missing Stack Sans Text face."),

 ("M-06", "MEDIUM", "Dual blog source of truth (Sanity + static blogData.ts) creates orphaned posts",
  "Code Quality",
  "blog/[slug] falls back from Sanity to src/data/blogData.ts, but the blog index reads Sanity only. Static-only posts "
  "are reachable by URL (and in the fallback sitemap) yet never listed; a ~300-line parallel static renderer diverges "
  "from the PortableText renderer.",
  "Inconsistent/unreachable content and duplicated rendering logic to maintain.",
  "Choose one source. If Sanity is canonical, remove blogData.ts and the static branch; otherwise surface static posts on the index."),

 ("M-07", "MEDIUM", "Dead Refine data provider points at a public demo API",
  "Code Quality",
  "src/providers/data-provider/index.ts uses API_URL = https://api.fake-rest.refine.dev; nothing in the app calls it.",
  "Confusing dead dependency shipped to clients; part of the Refine removal (H-12/C-01).",
  "Delete with the rest of the Refine stack."),

 ("M-08", "MEDIUM", "Images fetched as raw Sanity CDN URLs (image pipeline unused)",
  "Performance / Code Quality",
  "GROQ queries return mainImage.asset->url and render it directly; src/sanity/lib/image.ts (urlFor) is never imported.",
  "Sanity images ship without width/format/quality transforms — larger than necessary; dead helper code.",
  "Either delete image.ts or (better) render via urlFor(...).width(...).format('webp').url() for optimised delivery."),

 ("M-09", "MEDIUM", "Color-contrast failures (confirmed by Lighthouse)",
  "Accessibility",
  "Lighthouse flags color-contrast on the live site. Sources: muted/placeholder text (text-[#1E1E24]/20–60), inactive "
  "tab text (#A0A0A0 on white ≈ 2.6:1), footer gray-500 links, and white labels on the right (cyan) end of gradient "
  "CTA buttons.",
  "Informational text and controls fall below WCAG 1.4.3 (4.5:1), hurting low-vision users.",
  "Darken muted text to ≥ ~#595959; raise inactive tab to ~#767676; verify gradient-button label contrast and adjust."),

 ("M-10", "MEDIUM", "Mobile menu toggle lacks aria-expanded/aria-controls and focus handling",
  "Accessibility",
  "src/components/layout/Header.tsx hamburger has aria-label but no aria-expanded/aria-controls; opening the menu "
  "doesn’t move focus and Escape doesn’t close it.",
  "Screen-reader users can’t tell the menu state; keyboard UX is poor (WCAG 4.1.2).",
  "Add aria-expanded={open} + aria-controls=\"mobile-menu\"; manage focus and add Escape-to-close."),

 ("M-11", "MEDIUM", "Custom tab and single-select controls lack ARIA roles",
  "Accessibility",
  "WhoSection tabs (Leaders/Teams/Talent) use bare <button>s with no tablist/tab/tabpanel/aria-selected; the "
  "PrimaryClientContactForm timeline selector conveys selection by colour only (no aria-pressed / radiogroup).",
  "Assistive tech can’t convey which option is active (WCAG 1.3.1 / 4.1.2) and colour-only state fails 1.4.1.",
  "Add proper tablist/tab/tabpanel semantics; model the timeline as a radiogroup or add aria-pressed."),

 ("M-12", "MEDIUM", "Phone country-selector button has no accessible name",
  "Accessibility",
  "react-international-phone’s flag/country toggle (CandidateRegistrationForm, PrimaryClientContactForm, ApplyModal) "
  "renders an icon-only button; in ApplyModal the phone input itself isn’t programmatically labelled.",
  "Screen readers announce an unnamed “button” (WCAG 4.1.2).",
  "Pass aria-label to the country button and inputProps={{'aria-label':'Phone number'}} to the field."),

 ("M-13", "MEDIUM", "Favicon metadata points to a non-existent /public/favicon.ico",
  "SEO / Code Quality",
  "layout.tsx sets icons.icon:'/favicon.ico' but the file lives at src/app/favicon.ico; /public/favicon.ico does not "
  "exist, so the explicit metadata overrides the working App-Router convention with a 404 path.",
  "Possible missing/incorrect favicon in some contexts.",
  "Remove the explicit icons.icon (let the App-Router favicon.ico serve) or copy the file into /public."),

 ("M-14", "MEDIUM", "No preconnect + ~109 KB unused JS + legacy (ES5) JS on mobile",
  "Performance",
  "Lighthouse mobile diagnostics: uses-rel-preconnect (est. 320 ms), unused-javascript ~109 KiB, legacy-javascript "
  "~13 KiB (the ES5 target ships needless polyfills), plus render-blocking resources.",
  "Slower mobile start-up and wasted bytes.",
  "Add <link rel=preconnect> for cdn.sanity.io / GTM; remove Refine to cut unused JS (H-12); raise the TS target (M-18)."),

 ("M-15", "MEDIUM", "PrimaryClientContactForm field contract mismatches its API",
  "Code Quality",
  "The form sends the full international number in phoneNumber and an unused roleLabel; the API reads body.countryCode "
  "(always empty) and recomputes role labels server-side, so countryCode handling and the client roleLabel are dead.",
  "Fragile, inconsistent contract vs ApplyModal (which splits country code); confusing for maintainers.",
  "Standardise: either split countryCode like ApplyModal or drop it server-side; remove the unused roleLabel payload."),

 ("M-16", "MEDIUM", "No in-handler guard against double submission",
  "Code Quality / UX",
  "All four forms disable the submit button while loading but none early-returns inside handleSubmit when "
  "status===\"loading\"; a fast second Enter/submit can fire two POSTs (= duplicate lead emails).",
  "Duplicate submissions / duplicate notification + auto-reply emails.",
  "Add `if (status===\"loading\") return;` at the top of each handleSubmit."),

 ("M-17", "MEDIUM", "Inconsistent file/component naming and import aliases",
  "Code Quality",
  "Typo filenames SolutioinsBanner.tsx / SolutioinsAISpeed.tsx; SolutionSection exports lowercase Solutionsection; "
  "solutions/page.tsx uses deep relative imports while the rest use @components/*; tsconfig defines dead @pages/* and "
  "@sanity/* aliases.",
  "Hygiene/readability; raises onboarding friction and risk of mistakes.",
  "Rename to correct spelling + PascalCase; standardise on @components/*; prune unused aliases."),

 ("M-18", "MEDIUM", "tsconfig target is es5 (outdated for Next 16 / React 19)",
  "Code Quality / Performance",
  "tsconfig.json sets target:\"es5\". The code uses ES2021 APIs (String.replaceAll) and Node ≥20 is required; es5 is "
  "misleading and triggers legacy-JS downlevelling (see M-14).",
  "Unnecessary polyfills/larger output and tooling confusion.",
  "Set target to ES2022 (or esnext) to match the Next.js default."),

 ("M-19", "MEDIUM", "Deployment is stale relative to the codebase (~29 days behind)",
  "Operational",
  "Production is dated 2026-05-20 (sitemap lastmod; edge Age ≈29 days). Recent repo commits/working changes (address "
  "updates, navbar responsiveness, blog link marks, new blog upload scripts) are not yet live.",
  "Fixes already written aren’t benefiting users; audit findings should be re-verified after the next deploy.",
  "Establish a regular deploy cadence (or auto-deploy on merge to master); redeploy after applying audit fixes."),

 # ---------------- LOW ----------------
 ("L-01", "LOW", "Sitemap lastModified is always build time (no real freshness signal)",
  "SEO",
  "sitemap.ts sets every lastModified to new Date(); all entries share one timestamp that resets each build.",
  "No meaningful freshness signal to crawlers.",
  "Use Sanity _updatedAt/publishedAt for blog/job entries and fixed/real dates for static pages."),

 ("L-02", "LOW", "Generic / decorative alt text",
  "SEO / Accessibility",
  "Meaningful images use vague alts (“Hero Visual”, “AI Visual”, “Blog Hero Visual”); many decorative images carry "
  "describing alts (“Line”, “Background”, “Decoration”) instead of alt=\"\" + aria-hidden.",
  "Minor SEO/AT noise.",
  "Describe meaningful images with keywords; mark purely decorative ones alt=\"\" aria-hidden."),

 ("L-03", "LOW", "Verbose / inconsistent server logging",
  "Code Quality",
  "contact/route.ts logs full SMTP error name+message; the other three routes catch silently — inconsistent.",
  "Log hygiene only (client responses are correctly generic).",
  "Standardise on one logger across all routes."),

 ("L-04", "LOW", "External link in job page missing rel=\"noopener\"",
  "Security / Code Quality",
  "jobs/[slug] uses a raw <a> to talentifi-x.com without rel=noopener (blog external links do it correctly).",
  "Minor tab-nabbing / performance nit.",
  "Add rel=\"noopener noreferrer\" to external anchors."),

 ("L-05", "LOW", "not-found.tsx has no metadata",
  "SEO",
  "The 404 page has an <h1> but no metadata export (inherits the default title).",
  "Negligible (404s are noindex by status).",
  "Optionally add title:\"Page Not Found\"."),

 ("L-06", "LOW", "FAQ accordion missing aria-controls / panel id",
  "Accessibility",
  "blog/FaqSection uses a real button with aria-expanded but no aria-controls; the collapsed panel uses grid-rows[0fr] "
  "rather than hidden, so content may remain reachable when visually collapsed.",
  "Minor AT ambiguity.",
  "Add aria-controls + panel id; hide content when collapsed."),

 ("L-07", "LOW", "Decorative SVG/image aria-hidden inconsistency",
  "Accessibility",
  "Some inline SVGs (X logo) set aria-hidden while Instagram/LinkedIn don’t; a few decorative images use alt=\"\" without "
  "aria-hidden.",
  "Cosmetic; links are already labelled.",
  "Add aria-hidden=true + focusable=false to all decorative SVGs for consistency."),

 ("L-08", "LOW", "Validation ordering quirk in CandidateRegistrationForm",
  "Code Quality",
  "validateBeforeSubmit early-returns \"\" for a non-‘other’ heardFrom before later checks — harmless today but fragile "
  "if new checks are appended.",
  "Latent maintainability risk.",
  "Reorder so the empty-heardFrom check precedes the success return."),

 ("L-09", "LOW", "Twitter card has a handle but no creator/image",
  "SEO",
  "layout.tsx sets twitter.site:@talentifi_x with no twitter.creator and (until H-02) no image; cards degrade to text.",
  "Weaker X/Twitter previews.",
  "Verify the handle; add a twitter image (resolved by H-02)."),

 ("L-10", "LOW", "Minor source typos",
  "Code Quality",
  "SolutioinsAISpeed.tsx has a stray ‘x’ in a className and a commented-out sizes prop on an <Image>.",
  "Cosmetic; the commented sizes means that image ships without responsive hints.",
  "Remove the stray character; restore a sizes attribute."),

 ("L-11", "LOW", "robots/studio rely on disallow rather than a noindex directive",
  "SEO / Security",
  "/studio is disallowed in robots.txt but carries no robots:{index:false}; relies solely on the crawl hint + the "
  "hosting restriction.",
  "Low — defence in depth only.",
  "Add an explicit noindex to the studio route metadata."),

 ("L-12", "LOW", "Lighthouse render-blocking + preconnect already noted; misc diagnostics",
  "Performance",
  "Minor render-blocking resources and the absence of preconnect (also in M-14); GA could move to lazyOnload to "
  "shave mobile TBT.",
  "Small mobile start-up gains.",
  "Consider strategy=lazyOnload for analytics; add preconnects."),
]

POSITIVES = [
 "HTTPS enforced everywhere; all domain variants 301 to the canonical host; HSTS present.",
 "Security headers deployed in production: X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy.",
 "Email API routes are well built: SMTP header-injection guards (safeField), full HTML escaping, and (candidate route) "
 "file-upload size + MIME allow-listing.",
 "No dangerouslySetInnerHTML anywhere; PortableText text content auto-escapes.",
 "Excellent desktop performance (Lighthouse 97; LCP 1.0 s) and a perfect CLS of 0 on both profiles.",
 "Text compression (gzip/brotli) active; strong infra (Vercel edge, Mumbai region, ~0.3 s TTFB).",
 "Inter loaded via next/font (self-hosted, preloaded); accessible skip-link with a valid target; html lang set.",
 "Sensible per-page <title>/description metadata with a title template; responsive Tailwind layouts; keyboard-accessible "
 "file dropzone and Back-to-Top/Cookie buttons.",
]

# ============================================================
# RENDER CRITICAL SECTION (detailed)
# ============================================================
def render_finding(f, detailed=True):
    fid, sev, title, area, evidence, impact, fix = f
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    sev_run(p, sev)
    r = p.add_run(f"  {fid}  {title}")
    r.font.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = RGBColor.from_string(DARK)
    pa = para("", after=4, before=0)
    ra = pa.add_run("Area: "); ra.font.bold = True; ra.font.size = Pt(9.5); ra.font.color.rgb = RGBColor.from_string(GREY)
    ra2 = pa.add_run(area); ra2.font.size = Pt(9.5); ra2.font.color.rgb = RGBColor.from_string(GREY)
    for label, body in (("Evidence", evidence), ("Impact", impact), ("Recommendation", fix)):
        pp = doc.add_paragraph()
        pp.paragraph_format.space_after = Pt(3)
        rl = pp.add_run(f"{label}.  "); rl.font.bold = True; rl.font.size = Pt(10); rl.font.color.rgb = RGBColor.from_string(PRIMARY if label=="Recommendation" else DARK)
        rb = pp.add_run(body); rb.font.size = Pt(10)
    hr()

doc.add_heading("3. Critical Findings (Fix Immediately)", level=1)
para("These three issues should be remediated before other work; the first in particular undermines almost every "
     "other SEO investment on the site.", italic=True, color=GREY, after=8)
for f in [f for f in FINDINGS if f[1] == "CRITICAL"]:
    render_finding(f)
page_break()

# ============================================================
# FINDINGS BY CATEGORY
# ============================================================
doc.add_heading("4. Detailed Findings by Category", level=1)
para("Critical items are detailed in Section 3 and only summarised here. High findings are detailed in full; "
     "Medium and Low findings are listed with their recommended fix.", italic=True, color=GREY, after=6)

CATEGORY_ORDER = [
    ("SEO & Crawlability", ["SEO"]),
    ("Security & Compliance", ["Security"]),
    ("Performance", ["Performance"]),
    ("Accessibility", ["Accessibility"]),
    ("Code Quality & Architecture", ["Code Quality", "Operational", "Architecture"]),
]
sev_rank = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}

def in_cat(area, keys):
    return any(k.lower() in area.lower() for k in keys)

_claimed = set()
for catname, keys in CATEGORY_ORDER:
    # assign each finding to the FIRST category whose keywords match its area
    items = [f for f in FINDINGS if in_cat(f[3], keys) and f[0] not in _claimed]
    for f in items:
        _claimed.add(f[0])
    items = sorted(items, key=lambda f: (sev_rank[f[1]], f[0]))
    if not items:
        continue
    doc.add_heading(catname, level=2)
    # criticals are already detailed in Section 3 — reference them here, don't repeat
    crits = [f for f in items if f[1] == "CRITICAL"]
    for f in crits:
        p = para("", after=3)
        sev_run(p, "CRITICAL")
        r = p.add_run(f"  {f[0]}  {f[2]} — detailed in Section 3.")
        r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = RGBColor.from_string(GREY)
    # highs detailed in full
    for f in items:
        if f[1] == "HIGH":
            render_finding(f)
    # medium/low summary table
    ml = [f for f in items if f[1] in ("MEDIUM", "LOW")]
    if ml:
        para("Medium & Low findings", bold=True, size=10.5, color=PRIMARY, before=6, after=3)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for c, h in zip(t.rows[0].cells, ["ID", "Sev", "Finding", "Recommended fix"]):
            set_cell_text(c, h, bold=True, white=True, size=9)
            shade(c, PRIMARY)
        widths = [Inches(0.5), Inches(0.7), Inches(2.7), Inches(3.0)]
        for f in ml:
            cells = t.add_row().cells
            set_cell_text(cells[0], f[0], bold=True, size=8.5)
            set_cell_text(cells[1], f[1].title(), bold=True, size=8.5, color=SEV[f[1]])
            set_cell_text(cells[2], f[2], size=8.5)
            set_cell_text(cells[3], f[6], size=8.5)
            for c, w in zip(cells, widths):
                c.width = w
        for c, w in zip(t.rows[0].cells, widths):
            c.width = w
    doc.add_paragraph()

# de-dupe note: ensure each finding rendered once. Some areas list multiple keys; SEO/Accessibility overlaps
page_break()

# ============================================================
# PERFORMANCE DEEP-DIVE TABLE
# ============================================================
doc.add_heading("5. Core Web Vitals Detail (Lighthouse lab)", level=1)
t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
for c, h in zip(t.rows[0].cells, ["Metric", "Mobile", "Desktop", "Target"]):
    set_cell_text(c, h, bold=True, white=True, size=10); shade(c, PRIMARY)
metrics = [
    ("Performance score", "47", "97", "≥ 90"),
    ("First Contentful Paint", "4.0 s", "0.8 s", "< 1.8 s"),
    ("Largest Contentful Paint", "6.4 s", "1.0 s", "< 2.5 s"),
    ("Total Blocking Time", "720 ms", "0 ms", "< 200 ms"),
    ("Cumulative Layout Shift", "0", "0", "< 0.1"),
    ("Speed Index", "5.7 s", "0.9 s", "< 3.4 s"),
    ("Total transfer", "881 KiB", "1,698 KiB", "as low as possible"),
]
for m, mo, de, tg in metrics:
    cells = t.add_row().cells
    set_cell_text(cells[0], m, bold=True, size=9.5)
    set_cell_text(cells[1], mo, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
    set_cell_text(cells[2], de, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
    set_cell_text(cells[3], tg, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=GREY)
para("Interpretation: the codebase and infrastructure are capable of excellent performance (desktop proves it). The "
     "mobile shortfall is driven by (1) the client-side-rendering bailout delaying content paint, and (2) very heavy "
     "image sources. Fixing C-01 and H-07 should move mobile well into the green.",
     italic=True, color=GREY, before=6)
page_break()

# ============================================================
# WHAT'S WORKING WELL
# ============================================================
doc.add_heading("6. What’s Working Well", level=1)
para("A balanced audit notes strengths as well as gaps. The following are already done correctly and should be preserved:",
     after=4)
for s in POSITIVES:
    bullet(s, lead="✓", lead_color=SEV["POSITIVE"])
page_break()

# ============================================================
# ROADMAP
# ============================================================
doc.add_heading("7. Prioritised Remediation Roadmap", level=1)
para("A pragmatic sequence. Effort is a rough engineering estimate (S ≤ half-day, M ≤ 2 days, L > 2 days).",
     italic=True, color=GREY, after=6)

phases = [
 ("Phase 0 — This week (Critical)", PRIMARY, [
    ("C-01", "Remove Refine / restore server-side rendering for all pages", "M"),
    ("C-02", "Purge .env from git history; force-push; confirm key rotation", "S"),
    ("C-03", "Consolidate to one canonical host (www.talentifix.com) via env; fix sitemap, robots, blog canonicals", "M"),
 ]),
 ("Phase 1 — This sprint (High)", SECONDARY, [
    ("H-07", "Compress all images; trim the 8.5 MB hero; add preconnect", "M"),
    ("H-01/H-02/H-03", "Add metadataBase, default OG image, and JSON-LD (Org, JobPosting, Article, FAQ)", "M"),
    ("H-04", "Add jobs to the sitemap with final canonical URLs", "S"),
    ("H-05", "Rate limiting + honeypot/CAPTCHA on all form APIs; restrict Sanity CORS", "M"),
    ("H-06", "Gate GA/Clarity on cookie consent (Consent Mode v2)", "M"),
    ("H-08/H-09/H-10", "Modal focus trap; announce form errors; fix H1s / nested main / footer h1", "M"),
    ("H-11/H-12", "useCdn:true; remove embedded Studio from the marketing bundle", "S"),
 ]),
 ("Phase 2 — Following weeks (Medium)", "B7950B", [
    ("M-01", "Add CSP; extend HSTS (includeSubDomains; preload)", "M"),
    ("M-02", "Sanitise PortableText link/image URI schemes", "S"),
    ("M-04/M-05/M-08", "quality=75; WOFF2 + next/font; use Sanity image pipeline", "M"),
    ("M-06/M-07", "Resolve dual blog source; finish Refine removal", "M"),
    ("M-09–M-13", "Contrast fixes; menu/tab/phone ARIA; favicon path", "M"),
    ("M-15–M-19", "Form contract + double-submit guard; naming; TS target; deploy cadence", "M"),
 ]),
 ("Phase 3 — Opportunistic (Low)", GREY, [
    ("L-01–L-12", "Sitemap lastmod, alt text, logging, rel=noopener, 404 metadata, accordion ARIA, misc polish", "S"),
 ]),
]
for title, col, rows in phases:
    para(title, bold=True, size=11.5, color=col, before=8, after=3)
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
    for c, h in zip(t.rows[0].cells, ["Ref", "Action", "Effort"]):
        set_cell_text(c, h, bold=True, white=True, size=9); shade(c, col)
    ws = [Inches(1.0), Inches(5.0), Inches(0.6)]
    for ref, action, eff in rows:
        cells = t.add_row().cells
        set_cell_text(cells[0], ref, bold=True, size=9)
        set_cell_text(cells[1], action, size=9)
        set_cell_text(cells[2], eff, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9)
        for c, w in zip(cells, ws): c.width = w
    for c, w in zip(t.rows[0].cells, ws): c.width = w
page_break()

# ============================================================
# APPENDICES
# ============================================================
doc.add_heading("8. Appendices", level=1)

doc.add_heading("Appendix A — Domain & redirect map (live)", level=2)
t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
for c, h in zip(t.rows[0].cells, ["URL requested", "Result"]):
    set_cell_text(c, h, bold=True, white=True, size=9.5); shade(c, PRIMARY)
rmap = [
    ("https://www.talentifi-x.com/  (owner-stated canonical)", "301 → https://www.talentifix.com/"),
    ("https://talentifi-x.com/", "301 → https://www.talentifix.com/"),
    ("https://talentifix.com/", "301 → https://www.talentifix.com/"),
    ("https://www.talentifix.com/", "200 OK — the live site"),
    ("Sitemap entries (sitemap.xml)", "Listed as https://talentifix.com/… → all 301-redirect"),
    ("Blog canonical tags (source)", "Hard-coded https://www.talentifi-x.com/blog/… → redirects"),
]
for u, r in rmap:
    cells = t.add_row().cells
    set_cell_text(cells[0], u, size=9); set_cell_text(cells[1], r, size=9)

doc.add_heading("Appendix B — Largest image assets (top 12 of 103; 112.1 MB total)", level=2)
t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
for c, h in zip(t.rows[0].cells, ["File", "Size"]):
    set_cell_text(c, h, bold=True, white=True, size=9.5); shade(c, PRIMARY)
imgs = [
    ("public/banner-home/banner.webp (LCP hero)", "8.48 MB"),
    ("public/banner-home/Next.webp", "7.88 MB"),
    ("public/assets/contact/contact-form-bg.png", "7.88 MB"),
    ("public/blogs/b2.png", "7.36 MB"),
    ("public/banner-home/solution.webp", "7.21 MB"),
    ("public/assets/Solutions/perm-placement-bg.png", "7.21 MB"),
    ("public/blogs/b1.png", "6.82 MB"),
    ("public/blogs/b3.png", "6.72 MB"),
    ("public/assets/Solutions/hero-visual-blog.png", "6.00 MB"),
    ("public/assets/contact/contact-hero.png", "5.46 MB"),
    ("public/assets/figma/who-polygon.png", "4.78 MB"),
    ("public/assets/Solutions/ai-visual.svg", "2.21 MB"),
]
for f, s in imgs:
    cells = t.add_row().cells
    set_cell_text(cells[0], f, size=9); set_cell_text(cells[1], s, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("Appendix C — Tools & environment", level=2)
bullet("Live inspection: curl (HTTP headers, redirects, SSR HTML, sitemap/robots, asset sizes), Node 24 HTML parsing.")
bullet("Lighthouse v12 (headless Chrome) — mobile and desktop, 18 Jun 2026.")
bullet("Static review: full source tree, Next.js 16.2 / React 19.2 / Sanity 5.20 / Tailwind 4.2, branch karan.")
bullet("Stack confirmed in production: Server: Vercel; edge region bom1 (Mumbai); compression gzip/brotli.")
para("Note: PageSpeed/CrUX field data was unavailable (keyless API quota exhausted); analytics/GSC were excluded at "
     "the owner’s request. Figures are Lighthouse lab measurements.", italic=True, color=GREY, size=9, before=4)

# footer note
hr()
para("Prepared as an end-to-end technical audit of talentifix.com — 18 June 2026. "
     "Findings reflect the production deployment and source tree at the time of audit; re-verify after remediation deploys.",
     italic=True, color=GREY, size=8.5)

out = r"c:\Users\Karan\Desktop\Projects\Digitally next_projects\Talentifix\TalentiFi-X Website Audit Report - 2026-06-18.docx"
doc.save(out)
print("SAVED:", out)
